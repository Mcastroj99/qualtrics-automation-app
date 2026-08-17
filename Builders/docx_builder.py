import os
import sys

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if BASE_DIR not in sys.path:
    sys.path.insert(0, BASE_DIR)

try:
    from docx import Document
except ImportError:
    Document = None

from Core.models import SurveySchema, QuestionType, SelectorType


class DocxBuilder:
    """Genera el documento Word estandarizado 'Sample.docx' desde el esquema en memoria."""

    def build(self, schema: SurveySchema, output_file: str = "Sample.docx") -> str:
        if Document is None:
            # Fallback simple si python-docx no está instalado
            with open(output_file, "w", encoding="utf-8") as f:
                f.write(f"{schema.survey_name}\n\n")
                for i, q in enumerate(schema.questions, 1):
                    f.write(f"{i}. {q.question_text}\n")
            return output_file

        doc = Document()
        doc.add_heading(schema.survey_name or "Qualtrics Survey Specification", level=1)

        for i, q in enumerate(schema.questions, 1):
            if q.question_type == QuestionType.DB:
                doc.add_paragraph(q.question_text)
            elif q.question_type == QuestionType.TE and q.selector == SelectorType.FORM:
                p = doc.add_paragraph(f"{i}. {q.question_text} (Form field)")
                for c in q.choices:
                    doc.add_paragraph(f"* {c.text}")
            elif q.question_type == QuestionType.TE:
                doc.add_paragraph(f"{i}. {q.question_text} (Open text)")
            elif q.question_type == QuestionType.MATRIX:
                doc.add_paragraph(f"{i}. {q.question_text} (Matrix)")
                if q.answers:
                    doc.add_paragraph("Answer Choices:")
                    for a in q.answers:
                        doc.add_paragraph(a.text)
                if q.choices:
                    doc.add_paragraph("Statements:")
                    for c in q.choices:
                        doc.add_paragraph(c.text)
            elif q.question_type == QuestionType.SLIDER:
                doc.add_paragraph(f"{i}. {q.question_text} (Scale)")
                if q.answers:
                    scale_str = "Scale: " + " ".join([a.text for a in q.answers])
                    doc.add_paragraph(scale_str)
                if q.choices:
                    doc.add_paragraph("Statements:")
                    for c in q.choices:
                        doc.add_paragraph(c.text)
            elif q.is_multiple_answer or q.selector == SelectorType.MAVR:
                doc.add_paragraph(f"{i}. {q.question_text} (Multiple selection)")
                for c in q.choices:
                    doc.add_paragraph(c.text)
            else:
                doc.add_paragraph(f"{i}. {q.question_text} (Single selection)")
                for c in q.choices:
                    doc.add_paragraph(c.text)

            doc.add_paragraph("")  # Espaciado

        doc.save(output_file)
        print(f"✅ Documento Word maestro generado: {output_file}")
        return output_file