import os
from docx import Document
from pypdf import PdfReader


class DocumentParser:
    """Clase responsable de extraer el texto contenido en archivos DOCX, PDF o TXT."""

    @staticmethod
    def extract_text(file_path: str) -> str:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"El archivo no existe: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".docx":
            return DocumentParser._parse_docx(file_path)
        elif ext == ".pdf":
            return DocumentParser._parse_pdf(file_path)
        elif ext == ".txt":
            return DocumentParser._parse_txt(file_path)
        else:
            raise ValueError(f"Formato de archivo no soportado: {ext}")

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        doc = Document(file_path)
        full_text = []
        
        # Extracción de párrafos
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
                
        # Extracción de tablas (si la encuesta viene en formato tabular)
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    full_text.append(" | ".join(row_data))

        return "\n".join(full_text)

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
        return "\n".join(text)

    @staticmethod
    def _parse_txt(file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
